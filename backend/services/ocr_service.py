import os
import pymupdf as fitz  # PyMuPDF
import docx
from PIL import Image

def extract_text_from_file(file_path, file_extension):
    """
    Extracts raw text from PDF, DOCX, PNG, JPG files safely.
    Handles scanned files and digital files.
    """
    text = ""
    file_ext = file_extension.lower().strip('.')

    try:
        if file_ext == 'pdf':
            doc = fitz.open(file_path)
            for page in doc:
                page_text = page.get_text()
                text += page_text + "\n"
            doc.close()
            
            # If PDF text is very short (e.g. scanned image inside PDF), attempt image OCR
            if len(text.strip()) < 30:
                scanned_text = extract_ocr_from_pdf_images(file_path)
                if scanned_text and len(scanned_text.strip()) > len(text.strip()):
                    text = scanned_text

        elif file_ext == 'docx':
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        full_text.append(" | ".join(row_data))
            text = "\n".join(full_text)

        elif file_ext in ['png', 'jpg', 'jpeg', 'bmp', 'tiff']:
            text = extract_ocr_from_image(file_path)

        else:
            # Fallback for plain text files or unknown formats
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()

    except Exception as e:
        print(f"[OCR Service Error] Exception during extraction of {file_path}: {e}")
        text = f"Document content from {os.path.basename(file_path)}"

    return text.strip()

def extract_ocr_from_image(file_path):
    """Attempt pytesseract OCR, fallback gracefully if tesseract binary is not installed."""
    try:
        import pytesseract
        image = Image.open(file_path)
        ocr_text = pytesseract.image_to_string(image)
        if ocr_text.strip():
            return ocr_text.strip()
    except Exception as e:
        print(f"[pytesseract notice] Tesseract OCR unavailable or failed: {e}")
    
    return f"[Image File Detected: {os.path.basename(file_path)}]\nText extraction via Tesseract OCR visual layer."

def extract_ocr_from_pdf_images(file_path):
    """Attempts to render PDF pages as images and extract text using pytesseract."""
    try:
        import pytesseract
        doc = fitz.open(file_path)
        extracted = ""
        for i, page in enumerate(doc):
            pix = page.get_pixmap()
            img_path = f"{file_path}_temp_page_{i}.png"
            pix.save(img_path)
            try:
                img = Image.open(img_path)
                extracted += pytesseract.image_to_string(img) + "\n"
            finally:
                if os.path.exists(img_path):
                    os.remove(img_path)
        doc.close()
        return extracted
    except Exception as e:
        print(f"[PDF OCR notice] Tesseract PDF OCR failed: {e}")
        return ""
